"""
skills_service.py вЂ” 4 СЃРєРёР»Р»Р° Elira.

1. Р“РµРЅРµСЂР°С†РёСЏ Word/Excel
2. SQL Р·Р°РїСЂРѕСЃС‹ (SQLite)
3. HTTP/API РІС‹Р·РѕРІС‹
4. РЎРєСЂРёРЅС€РѕС‚ СЃР°Р№С‚Р° (playwright)
"""
from __future__ import annotations
import io
import json
import logging
import sqlite3
import time
from pathlib import Path
from typing import Any

import requests as http_lib

from app.core.config import DATA_DIR, GENERATED_DIR

logger = logging.getLogger(__name__)

OUTPUT_DIR = GENERATED_DIR
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def screenshot_capability_status() -> dict:
    try:
        from playwright.sync_api import sync_playwright  # noqa: F401
    except ImportError:
        return {
            "feature": "screenshot",
            "available": False,
            "reason": "optional_dependency_missing",
            "missing_packages": ["playwright"],
            "hint": "pip install playwright && playwright install chromium",
        }
    return {
        "feature": "screenshot",
        "available": True,
        "reason": None,
        "missing_packages": [],
        "hint": None,
    }


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 1. Р“Р•РќР•Р РђР¦РРЇ Р¤РђР™Р›РћР’
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def generate_word(title: str, content: str, filename: str = "") -> dict:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"ok": False, "error": "pip install python-docx"}

    doc = Document()
    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[0:3] in ("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. "):
            doc.add_paragraph(line[3:], style="List Number")
        else:
            doc.add_paragraph(line)

    fname = filename or f"elira_{int(time.time())}.docx"
    if not fname.endswith(".docx"):
        fname += ".docx"
    path = OUTPUT_DIR / fname
    doc.save(str(path))
    return {"ok": True, "path": str(path), "filename": fname, "size": path.stat().st_size,
            "download_url": f"/api/skills/download/{fname}"}


def generate_excel(title: str, data: list, headers: list = None, filename: str = "") -> dict:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return {"ok": False, "error": "pip install openpyxl"}

    wb = Workbook()
    ws = wb.active
    ws.title = title or "Sheet1"

    if headers:
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True, size=11)
            cell.fill = PatternFill(start_color="D5E8F0", end_color="D5E8F0", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

    start_row = 2 if headers else 1
    for row_idx, row_data in enumerate(data, start_row):
        if isinstance(row_data, (list, tuple)):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    fname = filename or f"elira_{int(time.time())}.xlsx"
    if not fname.endswith(".xlsx"):
        fname += ".xlsx"
    path = OUTPUT_DIR / fname
    wb.save(str(path))
    return {"ok": True, "path": str(path), "filename": fname, "size": path.stat().st_size,
            "download_url": f"/api/skills/download/{fname}"}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 2. SQL Р—РђРџР РћРЎР«
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

ALLOWED_DB_DIRS = [DATA_DIR.resolve()]

def _safe_db(db_path: str) -> Path:
    p = Path(db_path).resolve()
    for d in ALLOWED_DB_DIRS:
        try:
            p.relative_to(d)
            return p
        except ValueError:
            continue
    raise ValueError(f"Р—Р°РїСЂРµС‰РµРЅРѕ: {db_path}. РўРѕР»СЊРєРѕ data/")


def run_sql(db_path: str, query: str, params: list = None, max_rows: int = 100) -> dict:
    try:
        safe = _safe_db(db_path)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    if not safe.exists():
        return {"ok": False, "error": f"РќРµ РЅР°Р№РґРµРЅР°: {db_path}"}

    q_up = query.strip().upper()
    if any(q_up.startswith(c) for c in ["DROP ", "DELETE ", "TRUNCATE ", "ALTER "]):
        return {"ok": False, "error": f"Р—Р°Р±Р»РѕРєРёСЂРѕРІР°РЅРѕ: {q_up.split()[0]}"}

    try:
        conn = sqlite3.connect(safe)
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        if q_up.startswith(("SELECT", "PRAGMA", "WITH")):
            cur.execute(query, params or [])
            rows = cur.fetchmany(max_rows)
            columns = [d[0] for d in cur.description] if cur.description else []
            data = [dict(r) for r in rows]
            conn.close()
            return {"ok": True, "columns": columns, "rows": data, "count": len(data)}
        else:
            cur.execute(query, params or [])
            conn.commit()
            affected = cur.rowcount
            conn.close()
            return {"ok": True, "affected": affected}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def list_databases() -> dict:
    dbs = []
    for d in ALLOWED_DB_DIRS:
        if d.exists():
            for f in d.rglob("*.db"):
                dbs.append({"path": str(f), "name": f.name, "size": f.stat().st_size})
    return {"ok": True, "databases": dbs}


def describe_db(db_path: str) -> dict:
    try:
        safe = _safe_db(db_path)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    try:
        conn = sqlite3.connect(safe)
        tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name").fetchall()
        schema = {}
        for (tbl,) in tables:
            safe_tbl = '"' + tbl.replace('"', '""') + '"'
            cols = conn.execute(f"PRAGMA table_info({safe_tbl})").fetchall()
            cnt = conn.execute(f"SELECT COUNT(*) FROM {safe_tbl}").fetchone()[0]
            schema[tbl] = {"columns": [{"name": c[1], "type": c[2], "pk": bool(c[5])} for c in cols], "rows": cnt}
        conn.close()
        return {"ok": True, "tables": schema}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 3. HTTP / API
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

BLOCKED_HOSTS = {"localhost", "127.0.0.1", "0.0.0.0", "169.254.169.254"}

def http_request(url: str, method: str = "GET", headers: dict = None, body: Any = None, timeout: int = 15) -> dict:
    from urllib.parse import urlparse
    parsed = urlparse(url)
    if parsed.hostname in BLOCKED_HOSTS:
        return {"ok": False, "error": f"Р—Р°Р±Р»РѕРєРёСЂРѕРІР°РЅ: {parsed.hostname}"}

    try:
        kw = {"url": url, "headers": headers or {}, "timeout": timeout}
        method = method.upper()
        if method == "GET":
            resp = http_lib.get(**kw)
        elif method == "POST":
            kw["json"] = body if isinstance(body, (dict, list)) else None
            kw["data"] = body if not isinstance(body, (dict, list)) else None
            resp = http_lib.post(**kw)
        elif method == "PUT":
            kw["json"] = body if isinstance(body, (dict, list)) else None
            resp = http_lib.put(**kw)
        elif method == "DELETE":
            resp = http_lib.delete(**kw)
        else:
            return {"ok": False, "error": f"РќРµРёР·РІРµСЃС‚РЅС‹Р№ РјРµС‚РѕРґ: {method}"}

        ct = resp.headers.get("content-type", "")
        try:
            rbody = resp.json() if "json" in ct else resp.text[:30000]
        except Exception:
            rbody = resp.text[:30000]

        return {"ok": True, "status": resp.status_code, "body": rbody,
                "url": str(resp.url), "elapsed_ms": int(resp.elapsed.total_seconds() * 1000)}
    except http_lib.Timeout:
        return {"ok": False, "error": f"РўР°Р№РјР°СѓС‚ ({timeout}СЃ)"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ
# 4. РЎРљР РРќРЁРћРў РЎРђР™РўРђ
# в•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђв•ђ

def screenshot_url(url: str, width: int = 1280, height: int = 800, full_page: bool = False) -> dict:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        status = screenshot_capability_status()
        return {
            "ok": False,
            "error": "Screenshot feature is unavailable",
            "feature": status["feature"],
            "reason": status["reason"],
            "missing_packages": status["missing_packages"],
            "hint": status["hint"],
        }

    fname = f"screenshot_{int(time.time())}.png"
    path = OUTPUT_DIR / fname
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(viewport={"width": width, "height": height})
            page.goto(url, wait_until="networkidle", timeout=20000)
            page.screenshot(path=str(path), full_page=full_page)
            title = page.title()
            browser.close()
        return {"ok": True, "path": str(path), "filename": fname, "title": title,
                "download_url": f"/api/skills/download/{fname}", "view_url": f"/api/skills/view/{fname}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# 5. Task Planner Integration
# -----------------------------------------------------------------------------

def add_task(title: str, description: str = "", priority: str = "medium", due_date: str = None) -> dict:
    from app.services.task_planner_service import create_task
    return create_task(title, description, priority=priority, due_date=due_date)

def get_tasks(status: str = "todo") -> dict:
    from app.services.task_planner_service import list_tasks
    return list_tasks(status=status)

def set_task_status(task_id: str, status: str) -> dict:
    from app.services.task_planner_service import update_task
    return update_task(task_id, status=status)
