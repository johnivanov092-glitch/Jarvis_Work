"""
pdf_pro.py — продвинутая работа с PDF.

Возможности:
  • Извлечение текста (pypdf + pdfplumber fallback)
  • Таблицы (pdfplumber → структурированные данные)
  • OCR для сканированных PDF (pdf2image + pytesseract)
  • PDF → Word конвертация (сохраняет текст + таблицы)
  • Постраничный анализ

Зависимости (ставить по необходимости):
  pip install pypdf pdfplumber pytesseract pdf2image python-docx openpyxl
  
Для OCR также нужен Tesseract:
  Windows: https://github.com/UB-Mannheim/tesseract/wiki → установщик → добавить в PATH
  Для русского: скачать rus.traineddata в tessdata/
"""
from __future__ import annotations
import io
import logging
import time
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("data/generated")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════════════
# ИЗВЛЕЧЕНИЕ ТЕКСТА (умный: pypdf → pdfplumber → OCR)
# ═══════════════════════════════════════════════════════════════

def extract_pdf_smart(data: bytes, max_chars: int = 50000) -> dict:
    """
    Умное извлечение: пробует 3 метода по порядку.
    1. pypdf (быстро)
    2. pdfplumber (лучше с таблицами)
    3. OCR через pytesseract (для сканов)
    """
    results = {"text": "", "tables": [], "pages": 0, "method": "", "ocr_used": False}

    # Метод 1: pypdf (быстрый)
    text_pypdf = _try_pypdf(data, max_chars)
    if text_pypdf and len(text_pypdf.strip()) > 50:
        results["text"] = text_pypdf
        results["method"] = "pypdf"

    # Метод 2: pdfplumber (таблицы + лучший текст)
    text_plumber, tables, pages = _try_pdfplumber(data, max_chars)
    if text_plumber and len(text_plumber.strip()) > len(results["text"].strip()):
        results["text"] = text_plumber
        results["method"] = "pdfplumber"
    if tables:
        results["tables"] = tables
    results["pages"] = pages or _count_pages(data)

    # Метод 3: OCR (если текста мало — вероятно скан)
    if len(results["text"].strip()) < 100:
        ocr_text = _try_ocr(data, max_chars)
        if ocr_text and len(ocr_text.strip()) > 50:
            results["text"] = ocr_text
            results["method"] = "ocr"
            results["ocr_used"] = True

    return results


def _count_pages(data: bytes) -> int:
    try:
        from pypdf import PdfReader
        return len(PdfReader(io.BytesIO(data)).pages)
    except Exception:
        return 0


def _try_pypdf(data: bytes, max_chars: int) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts, total = [], 0
        for page in reader.pages:
            text = page.extract_text() or ""
            if total + len(text) > max_chars:
                parts.append(text[:max_chars - total])
                break
            parts.append(text)
            total += len(text)
        return "\n\n".join(parts)
    except Exception:
        return ""


def _try_pdfplumber(data: bytes, max_chars: int) -> tuple:
    """Возвращает (text, tables, page_count)."""
    try:
        import pdfplumber
    except ImportError:
        return "", [], 0

    try:
        pdf = pdfplumber.open(io.BytesIO(data))
        text_parts = []
        all_tables = []
        total = 0

        for i, page in enumerate(pdf.pages):
            # Текст
            page_text = page.extract_text() or ""
            if page_text.strip():
                if total + len(page_text) > max_chars:
                    break
                text_parts.append(f"--- Страница {i+1} ---\n{page_text}")
                total += len(page_text)

            # Таблицы
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if table and len(table) > 0:
                    all_tables.append({
                        "page": i + 1,
                        "index": t_idx,
                        "headers": table[0] if table[0] else [],
                        "rows": table[1:] if len(table) > 1 else [],
                        "row_count": len(table) - 1,
                    })

        page_count = len(pdf.pages)
        pdf.close()
        return "\n\n".join(text_parts), all_tables, page_count
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
        return "", [], 0


def _try_ocr(data: bytes, max_chars: int) -> str:
    """OCR через pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        return ""

    try:
        # Конвертируем PDF → изображения
        images = convert_from_bytes(data, dpi=200, first_page=1, last_page=10)  # Макс 10 страниц

        text_parts = []
        total = 0
        for i, img in enumerate(images):
            # OCR: пробуем русский + английский
            try:
                page_text = pytesseract.image_to_string(img, lang="rus+eng")
            except Exception:
                page_text = pytesseract.image_to_string(img, lang="eng")

            if page_text.strip():
                if total + len(page_text) > max_chars:
                    break
                text_parts.append(f"--- OCR страница {i+1} ---\n{page_text}")
                total += len(page_text)

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


# ═══════════════════════════════════════════════════════════════
# ТАБЛИЦЫ → Excel / CSV
# ═══════════════════════════════════════════════════════════════

def pdf_tables_to_excel(data: bytes, filename: str = "") -> dict:
    """Извлекает все таблицы из PDF и сохраняет в Excel."""
    try:
        import pdfplumber
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError:
        return {"ok": False, "error": "pip install pdfplumber openpyxl"}

    try:
        pdf = pdfplumber.open(io.BytesIO(data))
        wb = Workbook()
        wb.remove(wb.active)

        table_count = 0
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if not table or len(table) < 2:
                    continue
                table_count += 1
                ws = wb.create_sheet(title=f"P{i+1}_T{t_idx+1}")

                for r, row in enumerate(table, 1):
                    for c, val in enumerate(row or [], 1):
                        cell = ws.cell(row=r, column=c, value=val or "")
                        if r == 1:
                            cell.font = Font(bold=True)
                            cell.fill = PatternFill(start_color="D5E8F0", end_color="D5E8F0", fill_type="solid")

        pdf.close()

        if table_count == 0:
            return {"ok": False, "error": "Таблицы не найдены в PDF"}

        fname = filename or f"pdf_tables_{int(time.time())}.xlsx"
        if not fname.endswith(".xlsx"):
            fname += ".xlsx"
        path = OUTPUT_DIR / fname
        wb.save(str(path))

        return {"ok": True, "filename": fname, "tables": table_count, "size": path.stat().st_size,
                "download_url": f"/api/skills/download/{fname}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════
# PDF → WORD КОНВЕРТАЦИЯ
# ═══════════════════════════════════════════════════════════════

def pdf_to_word(data: bytes, filename: str = "") -> dict:
    """Конвертирует PDF → DOCX. Сохраняет текст, заголовки и таблицы."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"ok": False, "error": "pip install python-docx"}

    # Извлекаем контент
    result = extract_pdf_smart(data, max_chars=100000)
    text = result.get("text", "")
    tables = result.get("tables", [])

    if not text.strip() and not tables:
        return {"ok": False, "error": "Не удалось извлечь текст из PDF"}

    doc = Document()

    # Стиль
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Текст: парсим постранично
    current_page = ""
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("--- Страница") or stripped.startswith("--- OCR страница"):
            if current_page:
                doc.add_page_break()
            current_page = stripped
            continue
        if not stripped:
            continue

        # Угадываем заголовки (короткие, в верхнем регистре или с большим шрифтом)
        if len(stripped) < 80 and (stripped.isupper() or stripped.endswith(":")):
            doc.add_heading(stripped, level=2)
        else:
            doc.add_paragraph(stripped)

    # Таблицы
    for tbl in tables:
        headers = tbl.get("headers", [])
        rows = tbl.get("rows", [])
        if not headers and not rows:
            continue

        doc.add_paragraph("")  # Отступ
        doc.add_heading(f"Таблица (стр. {tbl.get('page', '?')})", level=3)

        all_rows = [headers] + rows if headers else rows
        max_cols = max(len(r) for r in all_rows) if all_rows else 0
        if max_cols == 0:
            continue

        table = doc.add_table(rows=len(all_rows), cols=max_cols)
        table.style = "Table Grid"

        for r, row in enumerate(all_rows):
            for c, val in enumerate(row[:max_cols]):
                cell = table.rows[r].cells[c]
                cell.text = str(val or "")
                if r == 0 and headers:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    fname = filename or f"pdf_converted_{int(time.time())}.docx"
    if not fname.endswith(".docx"):
        fname += ".docx"
    path = OUTPUT_DIR / fname
    doc.save(str(path))

    return {
        "ok": True,
        "filename": fname,
        "size": path.stat().st_size,
        "pages": result.get("pages", 0),
        "tables": len(tables),
        "method": result.get("method", ""),
        "ocr_used": result.get("ocr_used", False),
        "download_url": f"/api/skills/download/{fname}",
    }


# ═══════════════════════════════════════════════════════════════
# ПОСТРАНИЧНЫЙ АНАЛИЗ
# ═══════════════════════════════════════════════════════════════

def analyze_pdf(data: bytes) -> dict:
    """Подробный анализ PDF: страницы, текст, таблицы, изображения."""
    result = extract_pdf_smart(data)

    # Подсчёт слов
    words = len(result["text"].split()) if result["text"] else 0

    # Определяем тип PDF
    pdf_type = "text"
    if result.get("ocr_used"):
        pdf_type = "scanned"
    elif result.get("tables"):
        pdf_type = "tabular"
    elif words < 50:
        pdf_type = "image-heavy"

    return {
        "ok": True,
        "pages": result.get("pages", 0),
        "words": words,
        "chars": len(result.get("text", "")),
        "tables": len(result.get("tables", [])),
        "method": result.get("method", ""),
        "ocr_used": result.get("ocr_used", False),
        "type": pdf_type,
        "text_preview": result["text"][:500] if result["text"] else "",
        "table_summary": [
            {"page": t["page"], "rows": t["row_count"], "cols": len(t.get("headers", []))}
            for t in result.get("tables", [])[:10]
        ],
    }
